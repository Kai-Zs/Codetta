"""错题导出 DOCX"""
import io
import json
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from ..database import get_db

BODY_FONT = '微软雅黑'
CODE_FONT = 'Consolas'


def _set_font(run, name, size, bold=False):
    """设置字体，同时处理中西文字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _add_para(doc, text, font=BODY_FONT, size=11, bold=False, indent=None):
    """添加段落，统一字体"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_font(run, font, size, bold)
    return p


def export_wrong(user_id: int) -> bytes:
    with get_db() as db:
        rows = db.execute(
            "SELECT DISTINCT q.q_number, q.type, q.title, q.content, q.options, q.answer, "
            "q.answer_parts, q.template, q.answer_code, p.answer_status "
            "FROM progress p JOIN questions q ON p.question_id=q.id "
            "WHERE p.user_id=? AND p.answer_status IN ('incorrect','partial') AND p.removed_from_wrong=0 "
            "ORDER BY q.id",
            (user_id,)
        ).fetchall()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    for idx, r in enumerate(rows):
        if idx > 0:
            doc.add_paragraph("")
            doc.add_paragraph("")

        qtype = r["type"]

        # 题号 + 题型（小标题）
        _add_para(doc, f'第{r["q_number"]}题（{qtype}）', size=12, bold=True)

        # 题干内容（单选题截掉内嵌的选项文本）
        content = (r["content"] or "").strip()
        if qtype == "单选题" and content:
            m = re.search(r'\n([A-D][.、])', content)
            if m:
                content = content[:m.start()].strip()
        if content:
            _add_para(doc, content)

        # 选项（仅单选输出格式化选项）
        options = r["options"]
        if qtype == "单选题" and options:
            try:
                for opt in json.loads(options):
                    _add_para(doc, opt)
            except (json.JSONDecodeError, TypeError):
                pass

        # 填空答案提示（不重复"正确答案"）

        # 编程题预置代码
        template = r["template"]
        if template:
            _add_para(doc, "【编程题预置代码】", size=10, bold=True)
            _add_para(doc, template, font=CODE_FONT, size=9, indent=1)

        # 正确答案
        answer = r["answer"] or ""
        if qtype == "编程题":
            answer_text = (r["answer_code"] or "") or answer
            _add_para(doc, "正确答案：", size=11, bold=True)
            _add_para(doc, answer_text, font=CODE_FONT, size=9, indent=1)
        elif qtype == "填空题":
            try:
                parts = json.loads(r["answer_parts"] or "[]")
                answer_text = "  |  ".join(parts)
            except (json.JSONDecodeError, TypeError):
                answer_text = answer
            _add_para(doc, "正确答案：" + answer_text)
        else:
            _add_para(doc, "正确答案：" + answer)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
